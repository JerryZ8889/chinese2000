import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re, os

os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Read the markdown source
with open('Sophia中文识字进阶表_2000字.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Parse the markdown
stages = []
current_stage = None
current_cat = None

for line in md_content.split('\n'):
    stripped = line.strip()
    if stripped.startswith('## ') and '阶段' in stripped:
        current_stage = {'title': stripped[3:], 'categories': []}
        stages.append(current_stage)
        current_cat = None
    elif stripped.startswith('### ') and current_stage:
        current_cat = {'name': stripped[4:], 'chars': ''}
        current_stage['categories'].append(current_cat)
    elif stripped.startswith('#') or stripped.startswith('---') or stripped.startswith('>') or not stripped:
        continue
    elif current_cat:
        chars = re.findall(r'[\u4e00-\u9fff]', stripped)
        if chars:
            if current_cat['chars']:
                current_cat['chars'] += stripped
            else:
                current_cat['chars'] = stripped

# Verify
total_all = 0
for s in stages:
    total = sum(len(re.findall(r'[\u4e00-\u9fff]', c['chars'])) for c in s['categories'])
    total_all += total
    print(f"{s['title']}: {total} chars")
print(f"Total: {total_all}")


def add_styled_run(paragraph, text, bold=False, size=Pt(11), color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return run


def set_cell_content(cell, text, bold=False, size=Pt(11), color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    add_styled_run(p, text, bold=bold, size=size, color=color)


# Create docx
doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# === Title page ===
for _ in range(2):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_styled_run(p, 'Sophia 中文识字进阶表', bold=True, size=Pt(26), color=RGBColor(0x2B, 0x57, 0x97))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_styled_run(p, '2000字分阶段学习路线图', size=Pt(16), color=RGBColor(0x66, 0x66, 0x66))

doc.add_paragraph('')

summaries = [
    '第一阶段：1200字（基础识字）',
    '第二阶段：+300字（阅读拓展）→ 累计1500字',
    '第三阶段：+300字（表达深化）→ 累计1800字',
    '第四阶段：+200字（综合提升）→ 累计2000字',
]
for s in summaries:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(p, s, size=Pt(12))

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_styled_run(p, '配合《Sophia中文成长计划》使用', size=Pt(10), color=RGBColor(0x99, 0x99, 0x99))

doc.add_paragraph('')

# === Usage section ===
h = doc.add_heading('使用说明', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x97)

doc.add_paragraph('')
doc.add_paragraph(
    '本文档将 Sophia 中文学习计划中的 2000 个目标汉字按四个阶段分类呈现，'
    '每个阶段内按主题分组，方便查阅和跟踪学习进度。所有 2000 个汉字全局不重复。'
)

doc.add_paragraph('')
doc.add_heading('分阶段逻辑', level=2)
doc.add_paragraph(
    '第一阶段（1200字）：日常生活高频字 + Sophia兴趣领域（烘焙、画画、游戏）常用字，'
    '对应绘本和桥梁书阅读材料。'
)
doc.add_paragraph(
    '第二阶段（+300字）：章节书阅读中频繁遇到的字，包括叙事描写、心理活动、环境场景等方面。'
)
doc.add_paragraph(
    '第三阶段（+300字）：深度阅读和表达所需词汇，涵盖文学、文化、艺术、思维等领域。'
)
doc.add_paragraph(
    '第四阶段（+200字）：综合提升用字，包括抽象思辨、历史文化、成语构件和文学鉴赏。'
)

doc.add_paragraph('')
doc.add_heading('如何使用', level=2)
doc.add_paragraph('① 不需要按顺序一个一个学，这是查阅参考表，不是教材。结合日常阅读和练习中自然遇到的字来学。')
doc.add_paragraph('② 每个阶段末可以用对应的字表做识字测试：随机抽50个字，看Sophia能认出多少。')
doc.add_paragraph('③ 如果某个分类的字她大部分都认识，说明这个领域已经掌握得很好，可以跳过。')
doc.add_paragraph('④ 重点关注她不熟悉的分类，结合阅读和游戏中的实际场景来学习。')

doc.add_paragraph('')
doc.add_heading('注意事项', level=2)
doc.add_paragraph(
    '本表中 2000 个汉字严格去重，每个字仅出现一次。'
    '字表基于《现代汉语常用字表》和小学语文课标要求，'
    '并根据Sophia的实际场景（海外华裔、兴趣爱好）和分阶段阅读材料进行了调整和优化。'
)

doc.add_paragraph('')

# === Stage content with tables ===
stage_h1 = [
    '第一阶段：基础识字 1200 字',
    '第二阶段：阅读拓展 +300 字',
    '第三阶段：表达深化 +300 字',
    '第四阶段：综合提升 +200 字',
]

stage_sub = [
    '日常高频字 + Sophia兴趣领域常用字\n参考：《不一样的卡梅拉》《神奇校车》《米小圈上学记》',
    '章节书阅读中频繁遇到的中频字\n参考：继续《米小圈上学记》+ 更多章节书',
    '深度阅读和表达所需的进阶词汇\n参考：《笑猫日记》等复杂故事 + 中文视频讨论',
    '抽象概念与书面语综合提升\n参考：《夏洛的网》中文版 + 自选书籍 + 写作',
]

BORDER_XML = (
    '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>'
)

VALIGN_XML = '<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'
SHADE_XML = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F2F7FB" w:val="clear"/>'

for si, stage in enumerate(stages):
    # Stage heading
    h = doc.add_heading(stage_h1[si], level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2B, 0x57, 0x97)

    # Subtitle
    p = doc.add_paragraph(stage_sub[si])
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    # Tables for each category
    for cat in stage['categories']:
        m = re.match(r'(.+?)\s*\((\d+字)\)', cat['name'])
        if m:
            dname = m.group(1)
            dcount = m.group(2)
        else:
            dname = cat['name']
            dcount = ''

        all_chars = re.findall(r'[\u4e00-\u9fff]', cat['chars'])
        char_text = '\u3000'.join(all_chars)

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(
                '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            tbl.insert(0, tblPr)
        tblPr.append(parse_xml(BORDER_XML))

        table.columns[0].width = Cm(3.0)
        table.columns[1].width = Cm(14.0)

        # Left cell - category name with shading
        left = table.rows[0].cells[0]
        set_cell_content(
            left, f'{dname}\n({dcount})',
            bold=True, size=Pt(10),
            color=RGBColor(0x44, 0x72, 0xC4),
            align=WD_ALIGN_PARAGRAPH.CENTER
        )
        tc = left._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(VALIGN_XML))
        tcPr.append(parse_xml(SHADE_XML))

        # Right cell - characters
        right = table.rows[0].cells[1]
        set_cell_content(right, char_text, size=Pt(11))

        # Spacing paragraph
        sp = doc.add_paragraph('')
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(2)

    # Page break between stages (except last)
    if si < len(stages) - 1:
        doc.add_page_break()

# Save
output_path = '3 Sophia中文识字进阶表_2000字.docx'
doc.save(output_path)
print(f'\nSaved to: {output_path}')
print('Done!')
